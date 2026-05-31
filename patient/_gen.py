# -*- coding: utf-8 -*-
"""
Generates the 3 patent-application documents for the ArizaPro project,
modelled on the 3 reference templates the user received:
  1. muqova   — title page
  3. Model    — source-code listing of the program
  4. Referat  — abstract with metadata + functional description
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT = Path(__file__).resolve().parent

AUTHOR_FULL = "Umarov Kamoliddin Saloxiddin o'g'li"
AUTHORS = [
    "Umarov Kamoliddin Saloxiddin o'g'li",
    "Jo'raqulov Oybek Asadulla o'g'li",
    "Boybusinov Anvar Ural o'g'li",
]
AUTHORS_SHORT = ["Umarov K.S.", "Jo'raqulov O.A.", "Boybusinov A.U."]
COPYRIGHT_HOLDER = AUTHOR_FULL
CREATED_DATE = "18-may, 2026-yil"
YEAR = "2026"
CITY = "Toshkent"

PROGRAM_TITLE_UZ = (
    "ARIZA TURIDAGI HUJJATLARNI TELEGRAM MESSENJERI ORQALI "
    "AVTOMATIK TARZDA SHAKLLANTIRUVCHI DASTUR (ArizaPro)"
)
PROGRAM_TITLE_RU = (
    "(ПРОГРАММА ДЛЯ АВТОМАТИЧЕСКОГО ФОРМИРОВАНИЯ СУДЕБНЫХ "
    "ЗАЯВЛЕНИЙ ПОЛЬЗОВАТЕЛЕЙ ЧЕРЕЗ МЕССЕНДЖЕР TELEGRAM (ArizaPro))"
)


# ---------- helpers ----------

def set_default_font(doc, name="Times New Roman", size=14):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_margins(doc, top=2, bottom=2, left=3, right=1.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def p(doc, text, *, bold=False, size=None, align=None, italic=False, space_after=6):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    par.paragraph_format.space_after = Pt(space_after)
    return par


# ---------- 1. muqova ----------

def build_muqova():
    doc = Document()
    set_default_font(doc, "Times New Roman", 14)
    set_margins(doc)

    # vertical padding to push the title down a bit
    for _ in range(3):
        doc.add_paragraph("")

    p(doc, "EHM UCHUN DASTUR",
      bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    p(doc, PROGRAM_TITLE_UZ,
      bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    p(doc, PROGRAM_TITLE_RU,
      bold=True, size=14, italic=True,
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    p(doc, "HUQUQ EGASI",
      bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p(doc, COPYRIGHT_HOLDER,
      bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    p(doc, "Mualliflar:",
      bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    for a in AUTHORS:
        p(doc, a, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # push the year to the bottom
    for _ in range(6):
        doc.add_paragraph("")

    p(doc, f"{CITY} — {YEAR} yil",
      bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = OUT / "1.muqova_ArizaPro.docx"
    doc.save(out)
    print(f"Saved: {out.name}")


# ---------- 4. referat ----------

ANNOTATION_PARAS = [
    # vazifasi
    "Ushbu dasturiy ta'minot O'zbekiston Respublikasi fuqarolari uchun "
    "sud hujjatlarini (arizalar, shikoyatlar, iltimosnomalar va "
    "e'tiroznomalar) avtomatik tarzda shakllantirish maqsadida ishlab "
    "chiqilgan. Foydalanuvchi Telegram bot yoki Telegram Mini App "
    "(web-ilova) orqali savol-javob shaklida ma'lumot kiritadi va to'rt "
    "turdagi sud — fuqarolik, jinoyat, ma'muriy hamda iqtisodiy ishlar "
    "bo'yicha 30 dan ortiq tayyor andoza asosida rasmiy hujjat loyihasini "
    "tayyor PDF yoki DOCX holida oladi.",

    # qo'llanilish sohasi
    "Dasturning qo'llanilish sohasi — sud-huquqiy munosabatlar doirasidagi "
    "hujjatlarni elektron tarzda tayyorlash va sud jarayoni to'g'risidagi "
    "ma'lumotlarga ommaviy kirishni ta'minlash. Avvallari fuqarolar bunday "
    "hujjatlarni mustaqil tuza olmas, advokat yoki yuristga murojaat "
    "qilishga majbur edilar. ArizaPro buni bir necha daqiqaga qisqartiradi "
    "va xarajatni o'n barobardan ko'p kamaytiradi. Asosiy iste'molchilar — "
    "past va o'rta daromadli fuqarolar, mehnat migrantlari, yagona "
    "ota-onalar, shuningdek tadbirkorlik subyektlari.",

    # funksional imkoniyatlari (header)
    "Taklif etilayotgan dasturiy mahsulotning funksional imkoniyatlari:",

    "— foydalanuvchidan kerakli ma'lumotlarni bosqichma-bosqich (savol-javob "
    "shaklida) yig'ish; har bir maydon real vaqtda tasdiqlanadi (FISH, "
    "telefon raqami, sana, manzil, pul summasi, STIR/JSHSHIR va shu "
    "kabilar); shartli maydonlar oldingi javoblarga qarab ko'rsatiladi yoki "
    "yashiriladi;",

    "— bot bilan bir xil funksionalga ega zamonaviy Telegram Mini App "
    "(React asosidagi web-ilova): kim xohlasa bot orqali, kim xohlasa "
    "grafik interfeys orqali ishlaydi; initData HMAC-imzosi orqali "
    "foydalanuvchi autentifikatsiyasi amalga oshiriladi;",

    "— jadval2.sud.uz axborot tizimi bilan integratsiya: sud majlislariga "
    "tayinlangan ishlar ro'yxatini hudud va sud kesimida (jinoyat ishlari "
    "hamda ma'muriy huquqbuzarliklar bo'yicha) ko'rish va F.I.SH. yoki ish "
    "raqami bo'yicha qidirish (\"Ishimni tekshirish\");",

    "— O'zbekiston sudlari bo'yicha ma'lumotnoma: har bir sudning yuridik "
    "manzili, telefon raqami, elektron pochtasi va xaritadagi joylashuvi "
    "(\"Sudlar ma'lumoti\");",

    "— Telegram profilining ma'lumotlari asosida F.I.SH.ni avtomatik "
    "to'ldirish, tanlangan sud va hududga qarab sud nomi va sudya "
    "ma'lumotlarini avtomatik kiritish; hujjat tilini foydalanuvchining til "
    "afzalligiga qarab tanlash (o'zbek kirill, o'zbek lotin yoki rus tilida);",

    "— belgilangan andoza asosida DOCX hujjatini shakllantirish "
    "(docxtemplater texnologiyasi orqali) va undan keyin LibreOffice "
    "vositasida PDF formatga avtomatik o'tkazish;",

    "— erkin matnli maydonlarni ovozli xabar orqali to'ldirish: ovoz "
    "matnga aylantiriladi (Whisper/Groq nutqni tanish xizmati);",

    "— sun'iy intellekt yordamida foydalanuvchining erkin matnli "
    "javoblarini (masalan, e'tiroz sabablarini) rasmiy-yuridik uslubda "
    "qayta yozish (Anthropic Claude API integratsiyasi);",

    "— Click UZ va Payme (Paycom) to'lov tizimlari bilan integratsiya: "
    "har bir tayyor hujjat uchun belgilangan summada to'lov qabul qilish "
    "(JSON-RPC, MD5 imzo tekshiruvi, webhook orqali ma'lumotni qaytarish);",

    "— Telegram chati orqali QR-kod va deep-link orqali hujjatni qayta "
    "yuklab olish imkoniyati (har bir hujjatga noyob downloadToken "
    "biriktiriladi; eskirgan fayllar avtomatik tozalanib turiladi);",

    "— foydalanuvchilardan arizalar tuzilishi yuzasidan takliflar va "
    "shikoyatlarni qabul qilib, ularni xizmat ko'rsatuvchi guruhga "
    "yo'naltirish;",

    "— admin paneli orqali statistika, to'lov hisobotlari (Excel formatda) "
    "va barcha foydalanuvchilarga ommaviy xabar tarqatish;",

    "— oraliq holatni PostgreSQL ma'lumotlar bazasida saqlash, bu esa "
    "foydalanuvchi botni qayta ishga tushirsa, to'ldirilgan ma'lumotlar "
    "yo'qolmasligini ta'minlaydi.",

    "Dastur server qismi Node.js 20 platformasi, TypeScript dasturlash tili "
    "va Telegraf freymvorki asosida ishlab chiqilgan. Telegram Mini App "
    "(web-ilova) React 19, Vite va Tailwind CSS asosida tuzilgan. Ma'lumotlar "
    "bazasi sifatida PostgreSQL 17 va Prisma ORM ishlatiladi. Hujjat "
    "shakllantirish uchun docxtemplater + PizZip, PDF konvertatsiyasi uchun "
    "LibreOffice headless rejimi qo'llaniladi. Foydalanuvchi interfeysi "
    "3 tilda: o'zbek (kirill va lotin) va rus tilida.",
]


def build_referat():
    doc = Document()
    set_default_font(doc, "Times New Roman", 14)
    set_margins(doc)

    p(doc, "Referat",
      bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    p(doc, "Mualliflar: " + ", ".join(AUTHORS), space_after=10)
    p(doc, f"Huquqiy egasi: {COPYRIGHT_HOLDER}", space_after=10)

    p(doc, "EHM uchun dastur (Ma'lumotlar bazasi):", bold=True, space_after=4)
    p(doc, PROGRAM_TITLE_UZ, space_after=4)
    p(doc, PROGRAM_TITLE_RU, italic=True, space_after=12)

    p(doc, f"Tuzilgan vaqti: {CREATED_DATE}", space_after=12)

    p(doc,
      "Annotatsiya (vazifasi, qo'llanilish sohasi, funksional "
      "imkoniyatlari ko'rsatilsin):",
      bold=True, space_after=10)

    for para in ANNOTATION_PARAS:
        p(doc, para, space_after=8)

    # tech specs
    p(doc, "")
    p(doc, "EHM turi: Intel Core i3 (1.6 GGs) va undan yuqori; "
      "operativ xotira: 2 GB va undan yuqori")
    p(doc, "Dastur tili: TypeScript 5 (server — Node.js 20 LTS; "
      "web-ilova — React 19)")
    p(doc, "Server kutubxonalari: Telegraf 4.16, Prisma 5, docxtemplater, "
      "PizZip, ExcelJS, Sharp, qrcode, libphonenumber-js, Pino, Zod, "
      "@anthropic-ai/sdk")
    p(doc, "Telegram Mini App (web-ilova): React 19, Vite 6, "
      "Tailwind CSS 4, React Router 7")
    p(doc, "Ma'lumotlar bazasi: PostgreSQL 17")
    p(doc, "Tashqi xizmatlar va integratsiyalar: jadval2.sud.uz "
      "(sud majlislari jadvali), Click UZ va Payme to'lov tizimlari, "
      "Anthropic Claude API (matnni qayta yozish), Whisper/Groq "
      "(nutqni matnga aylantirish)")
    p(doc, "Operatsiya tizimi: Linux (Ubuntu 22.04+ tavsiya etiladi), "
      "Windows 10/11, macOS 12+")
    p(doc, "Tashqi vositalar: LibreOffice (DOCX → PDF konvertatsiyasi uchun)")
    p(doc, "Dastur hajmi (kompilyatsiyadan oldingi manba kod): ~700 KB "
      "(~20 000 satr)")

    out = OUT / "4.Referat_ArizaPro.docx"
    doc.save(out)
    print(f"Saved: {out.name}")


# ---------- 3. Model (source listing) ----------

LISTING_FILES = [
    # ── Bot core ──
    "src/main.ts",
    "src/bot/index.ts",
    "src/bot/commands.ts",
    "src/bot/actions.ts",
    "src/bot/keyboards.ts",
    "src/scenes/ariza-wizard.scene.ts",
    # ── Document generation ──
    "src/services/document.service.ts",
    "src/services/docx.service.ts",
    "src/services/pdf.service.ts",
    # ── Court schedule (jadval2) + court directory ──
    "src/services/jadval2.service.ts",
    "src/templates/court-info.ts",
    # ── AI rewrite + voice-to-text ──
    "src/services/ai-assist.service.ts",
    "src/services/transcription.service.ts",
    # ── Payments + Telegram Mini App backend ──
    "src/services/payment.service.ts",
    "src/services/webapp-api.ts",
    "src/services/telegram-auth.ts",
    # ── Maintenance ──
    "src/services/cleanup.service.ts",
    # ── Templates registry + data model ──
    "src/templates/registry.ts",
    "prisma/schema.prisma",
    # ── Telegram Mini App (React front-end) ──
    "webapp/src/App.tsx",
    "webapp/src/api.ts",
    "webapp/src/components/FieldEditor.tsx",
]


def add_code_block(doc, text, size=9):
    """Adds a paragraph in Courier New, preserving line breaks."""
    # python-docx doesn't keep multi-line text inside a single run, so we
    # split on \n and create one paragraph per source line — keeps tab/spaces
    # and avoids the auto-formatting Word does on long single-run paragraphs.
    style = doc.styles["Normal"]
    for line in text.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(size)
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Courier New")
        rfonts.set(qn("w:hAnsi"), "Courier New")
        rfonts.set(qn("w:cs"), "Courier New")
        rpr.append(rfonts)


def build_model():
    doc = Document()
    set_default_font(doc, "Times New Roman", 12)
    set_margins(doc, top=1.5, bottom=1.5, left=2, right=1.5)

    p(doc, PROGRAM_TITLE_UZ,
      bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p(doc, "Dastur manba kodining ro'yxati (listing of source code)",
      italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    for rel in LISTING_FILES:
        full = ROOT / rel
        if not full.exists():
            print(f"  ! missing: {rel}")
            continue
        # File header
        p(doc, f"// ============================================",
          size=11)
        p(doc, f"// File: {rel}", bold=True, size=12)
        p(doc, f"// ============================================",
          size=11, space_after=4)
        try:
            text = full.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = full.read_text(encoding="cp1251", errors="replace")
        add_code_block(doc, text, size=9)
        doc.add_paragraph("")  # spacer between files

    out = OUT / "3.Model_ArizaPro.docx"
    doc.save(out)
    print(f"Saved: {out.name}")


if __name__ == "__main__":
    build_muqova()
    build_referat()
    build_model()
    print("\nAll 3 patent documents generated in:", OUT)
